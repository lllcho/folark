"""DOCX 文件处理实现。"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING

from app.plugins.builtin_plugin.utils import render_template
from app.plugins.core import PreviewResult

if TYPE_CHECKING:
    from app.plugins.core import PluginContext

logger = logging.getLogger(__name__)


def extract(ctx: PluginContext) -> str | None:
    """从 DOCX 文件提取文本。"""
    try:
        from docx import Document

        doc = Document(ctx.file_path)
        text_parts = []

        # 遍历段落
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)

        # 遍历表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)

        text = "\n".join(text_parts).strip()
        return text if text else None
    except Exception as e:
        logger.warning("DOCX 文本提取失败 %s: %s", ctx.file_path, e)
        return None


def convert(ctx: PluginContext) -> bytes | None:
    """将 DOCX 文件转换为 PDF。"""
    if ctx.target_type != "pdf":
        logger.warning("DOCX 不支持转换为 %s 格式", ctx.target_type)
        return None

    try:
        import pymupdf

        doc = pymupdf.open(ctx.file_path)
        pdf_bytes = doc.convert_to_pdf()
        doc.close()

        return pdf_bytes
    except Exception as e:
        logger.warning("DOCX 转 PDF 失败 %s: %s", ctx.file_path, e)
        return None


def preview(ctx: PluginContext) -> PreviewResult:
    """使用模板渲染 DOCX 预览页面。"""
    html = render_template(
        "docx_preview.html",
        {
            "file_name": ctx.file_name,
            "title": ctx.title or ctx.file_name,
            "file_url": ctx.file_url,
        },
    )
    return PreviewResult.from_html(html)
